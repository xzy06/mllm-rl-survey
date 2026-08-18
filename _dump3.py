# -*- coding: utf-8 -*-
"""提取二稿.docx 全文（段落+表格），直接写入 UTF-8 文本文件"""
from docx import Document

doc = Document('二稿.docx')
lines = []
lines.append('===== 段落全文 =====')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        lines.append(f'[{i}] {t}')
lines.append('')
lines.append('===== 表格 =====')
for ti, tbl in enumerate(doc.tables):
    lines.append(f'--- 表{ti+1} ---')
    for row in tbl.rows:
        cells = [c.text.strip().replace('\n', '⏎') for c in row.cells]
        lines.append(' | '.join(cells))

with open('_draft3.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))
print('written', len(lines), 'lines')
