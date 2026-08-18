# -*- coding: utf-8 -*-
"""二稿.docx 第三章 B/C 类文献篇幅压缩（加码版，含公式处理）"""
from docx import Document
from docx.oxml.ns import qn

DOC = '二稿.docx'
doc = Document(DOC)

def replace_cross_runs(p, old, new):
    """段落级文本替换，支持 old 跨多个 w:t；非文本元素（oMath 等）位置不动"""
    ts = [t for t in p._p.iter(qn('w:t'))]
    texts = [t.text or '' for t in ts]
    full = ''.join(texts)
    if old not in full:
        return False
    pos = full.index(old)
    end = pos + len(old)
    starts = []
    acc = 0
    for t in ts:
        starts.append(acc)
        acc += len(t.text or '')
    si = 0
    for k, s in enumerate(starts):
        if s <= pos:
            si = k
    ei = si
    for k in range(si, len(ts)):
        if starts[k] + len(texts[k]) >= end:
            ei = k
            break
    if si == ei:
        t = ts[si]
        t.text = (t.text or '')[:pos - starts[si]] + new + (t.text or '')[end - starts[si]:]
    else:
        ts[si].text = (ts[si].text or '')[:pos - starts[si]] + new
        for k in range(si + 1, ei):
            ts[k].text = ''
        ts[ei].text = (ts[ei].text or '')[end - starts[ei]:]
    return True

def para_full(p):
    parts = []
    for child in p._p.iter():
        if child.tag == qn('w:t'):
            parts.append(child.text or '')
        elif child.tag == qn('m:t'):
            parts.append('⟦' + (child.text or '') + '⟧')
    return ''.join(parts)

# ---------- 1. [34] SpatialThinker 解释段：删权重句（含 4 公式）+ CIoU 稠密梯度句 ----------
p34 = doc.paragraphs[34]._p
r1_done = r5_done = False
for r in p34.findall(qn('w:r')):
    for t in r.findall(qn('w:t')):
        if '权重固定为' in (t.text or ''):
            t.text = (t.text or '').replace('；权重固定为', '。')
            r1_done = True
        if '准确率优先' in (t.text or ''):
            t.text = (t.text or '').replace('，准确率优先。', '')
            r5_done = True
        if '对不重叠的框也提供稠密梯度' in (t.text or ''):
            t.text = (t.text or '').replace('，对不重叠的框也提供稠密梯度', '')
# 删除 4 个权重内联公式与“、”“ ”连接 run
for ch in list(p34):
    if ch.tag == qn('m:oMath'):
        p34.remove(ch)
    elif ch.tag == qn('w:r'):
        texts = ''.join(t.text or '' for t in ch.iter(qn('w:t')))
        if texts in ('、', ' '):
            p34.remove(ch)
assert r1_done and r5_done, '34 段替换失败'

# ---------- 2. [35] SVQA-R1 段：删“验证器完全自监督：” ----------
assert replace_cross_runs(doc.paragraphs[35], '验证器完全自监督：', ''), '35 段替换失败'

# ---------- 3. [39] Ground-R1 段：机制句压缩 ----------
assert replace_cross_runs(doc.paragraphs[39],
    '答案信号会倒逼定位行为涌现：错误的定位会降低裁剪区域的答案正确率，从而压低对应轨迹的组内优势，定位行为与答案正确性在信用分配上被耦合。',
    '答案信号会倒逼定位行为涌现。'), '39 段替换失败'

# ---------- 4. [46] VisualPRM 句：删数据量 ----------
assert replace_cross_runs(doc.paragraphs[46], '以约 40 万条过程监督数据', ''), '46 段替换失败'

# ---------- 5. [52] H-GRPO 公式解释压缩（保留 3 个公式对象） ----------
p52 = doc.paragraphs[52]._p
for t in p52.iter(qn('w:t')):
    s = t.text or ''
    if '衡量预测证据框存在且与参考证据区域兼容' in s:
        t.text = s.replace(' 衡量预测证据框存在且与参考证据区域兼容，', '、')
    elif s.strip() == '和' and len(s) <= 3:
        t.text = '、'
    elif s.strip() == '' and len(s) <= 2:
        t.text = ''
    elif '为子问题与子答案的语义相似度' in s:
        t.text = s.replace(' 为子问题与子答案的语义相似度（Sentence-BERT 余弦），IoU 为两框的空间重叠。',
                           ' 分别衡量证据框兼容性、子问题与子答案的语义相似度，IoU 为空间重叠。')

# ---------- 6. [54] H-GRPO 段尾：删 RoboSpatial 数据 ----------
assert replace_cross_runs(doc.paragraphs[54], '、OOD 基准 RoboSpatial 达 70.2%', ''), '54 段替换失败'

doc.save(DOC)
print('saved')

# ---------- 验证 ----------
doc2 = Document(DOC)
for i in (34, 35, 39, 46, 52, 54):
    print(f'[{i}] {para_full(doc2.paragraphs[i])}')
