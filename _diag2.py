# -*- coding: utf-8 -*-
"""dump 目标段落完整元素结构（w:t 文本 + m:oMath 公式），供压缩脚本设计"""
from docx import Document
from docx.oxml.ns import qn

doc = Document('二稿.docx')
out = []

def dump_para(i):
    p = doc.paragraphs[i]
    out.append(f'===== [{i}] =====')
    for j, child in enumerate(p._p):
        tag = child.tag.split('}')[-1]
        if child.tag == qn('w:t'):
            out.append(f'  [{j}] w:t  文本={child.text!r}')
        elif child.tag == qn('m:oMath'):
            texts = ''.join(t.text or '' for t in child.iter(qn('m:t')))
            out.append(f'  [{j}] oMath 公式={texts!r}')
        elif tag in ('pPr', 'rPr'):
            out.append(f'  [{j}] {tag}')
        else:
            out.append(f'  [{j}] {tag}')

for i in (32, 34, 35, 39, 40, 41, 43, 46, 47, 49, 50, 52, 54):
    dump_para(i)
    p = doc.paragraphs[i]
    out.append(f'-- [{i}] w:t 列表 --')
    for k, t in enumerate(p._p.iter(qn('w:t'))):
        out.append(f'  t{k}: {t.text!r}')

with open('_diag2.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(out))
print('written, lines:', len(out))
