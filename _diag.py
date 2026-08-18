# -*- coding: utf-8 -*-
"""诊断目标段落元素分布：w:t 文本与 m:oMath 公式的排列，确认文本片段是否跨元素"""
from docx import Document
from docx.oxml.ns import qn

doc = Document('二稿.docx')

def dump_para(i):
    p = doc.paragraphs[i]
    out.append(f'===== [{i}] =====')
    for j, child in enumerate(p._p):
        tag = child.tag.split('}')[-1]
        if child.tag == qn('w:t'):
            out.append(f'  [{j}] w:t  文本={child.text!r}')
        elif child.tag == qn('m:oMath'):
            texts = ''.join(t.text or '' for t in child.iter(qn('m:t')))
            out.append(f'  [{j}] oMath 公式={texts!r}')
        elif tag in ('pPr', 'rPr'):
            out.append(f'  [{j}] {tag}')
        else:
            out.append(f'  [{j}] {tag}')

out = []
for i in (34, 35, 39, 46, 52, 54):
    dump_para(i)
    # 追加每个 w:t 的文本（含嵌套）
    p = doc.paragraphs[i]
    out.append(f'-- [{i}] w:t 列表 --')
    for k, t in enumerate(p._p.iter(qn('w:t'))):
        out.append(f'  t{k}: {t.text!r}')

with open('_diag.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(out))
print('written')
