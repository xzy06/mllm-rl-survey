# -*- coding: utf-8 -*-
"""Build v2 of 二稿_英文版.docx: MT text + mechanical format fixes only.
Fixes: math-adjacent spacing (restored per original zh chunk whitespace),
bold labels (Keywords / Section summary), space-before-punct, broken
superscript spacing (CR ³), proper-noun case (Polia/Mme/THE METHOD),
arrow spacing normalization. No re-translation, no wording changes."""
import json, re, sys
from copy import deepcopy
sys.path.insert(0, r'C:\Users\Lenovo\Desktop\study\essay\mllm-rl-survey\tmp')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
SRC = r"C:\Users\Lenovo\Desktop\study\essay\mllm-rl-survey\二稿.docx"
DST = r"C:\Users\Lenovo\Desktop\study\essay\mllm-rl-survey\二稿_英文版.docx"
TRANS = r"C:\Users\Lenovo\Desktop\study\essay\mllm-rl-survey\tmp\translations.json"

# ---------- table beautification ----------
HEADER_FILL = 'D9E2F3'   # light blue
BORDER_COLOR = '808080'

def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn('w:' + k), v)
    return e

def cell_border(cell, color=BORDER_COLOR, sz='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = tcPr.find(qn('w:tcBorders'))
    if tcB is None:
        tcB = OxmlElement('w:tcBorders'); tcPr.append(tcB)
    for edge in ('top', 'left', 'bottom', 'right'):
        ex = tcB.find(qn('w:' + edge))
        if ex is None:
            ex = OxmlElement('w:' + edge); tcB.append(ex)
        ex.set(qn('w:val'), 'single'); ex.set(qn('w:sz'), sz)
        ex.set(qn('w:space'), '0'); ex.set(qn('w:color'), color)

def cell_shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)

def cell_valign(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    va = tcPr.find(qn('w:vAlign'))
    if va is None:
        va = OxmlElement('w:vAlign'); tcPr.append(va)
    va.set(qn('w:val'), val)

def cell_runs_bold(cell, bold=True, size_pt=None):
    for p in cell.paragraphs:
        for r in p.runs:
            rPr = r._r.get_or_add_rPr()
            b = rPr.find(qn('w:b'))
            if bold and b is None:
                rPr.append(OxmlElement('w:b'))
            if size_pt is not None:
                sz = rPr.find(qn('w:sz'))
                if sz is None:
                    sz = OxmlElement('w:sz'); rPr.append(sz)
                sz.set(qn('w:val'), str(int(size_pt * 2)))

def cell_align(cell, align):
    for p in cell.paragraphs:
        p.alignment = align

def beautify_tables(doc):
    for tbl in doc.tables:
        # table centered + autofit to page width
        tblPr = tbl._tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            jc = tblPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc'); tblPr.append(jc)
            jc.set(qn('w:val'), 'center')
        rows = tbl.rows
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row.cells):
                cell_border(cell)
                cell_valign(cell, 'center')
                if ri == 0:                      # header row
                    cell_shade(cell, HEADER_FILL)
                    cell_runs_bold(cell, bold=True, size_pt=10.5)
                    cell_align(cell, WD_ALIGN_PARAGRAPH.CENTER)
                elif ci == 0:                    # first column (method/benchmark names)
                    cell_runs_bold(cell, bold=True)
                    cell_align(cell, WD_ALIGN_PARAGRAPH.LEFT)
                else:
                    cell_align(cell, WD_ALIGN_PARAGRAPH.LEFT)



def has_han(s):
    return any('\u4e00' <= c <= '\u9fff' for c in s)

def mech_fix(en):
    en = re.sub(r' ([\.,;:])', r'\1', en)          # space before punctuation
    en = en.replace('CR ³', 'CR³')                  # broken superscript spacing
    en = en.replace('Q-Spatial + +', 'Q-Spatial++')
    en = re.sub(r'\bV \*', 'V*', en)
    en = re.sub(r'\bMme\b', 'MME', en)              # proper-noun case
    en = re.sub(r'\bPolia\b', 'POLIA', en)
    en = en.replace('THE METHOD', 'Method')         # all-caps table header
    en = re.sub(r'\s*→\s*', ' → ', en)              # arrow spacing
    return en

def split_bold_label(p_el, label):
    """Make `label` its own bold run at paragraph start (formatting parity with 二稿)."""
    for t in p_el.iter(qn('w:t')):
        if t.text and t.text.startswith(label):
            run = t.getparent()
            rest = t.text[len(label):]
            new_run = deepcopy(run)
            rPr = new_run.find(qn('w:rPr'))
            if rPr is None:
                rPr = new_run.makeelement(qn('w:rPr'), {})
                new_run.insert(0, rPr)
            if rPr.find(qn('w:b')) is None:
                rPr.append(rPr.makeelement(qn('w:b'), {}))
            for tt in new_run.findall(qn('w:t')):
                tt.text = label
                tt.set(qn('xml:space'), 'preserve')
            t.text = rest
            t.set(qn('xml:space'), 'preserve')
            run.addprevious(new_run)
            return True
    return False

def main():
    trans = json.load(open(TRANS, encoding='utf-8'))
    doc = Document(SRC)
    body = doc.element.body
    applied = 0
    for p_el in body.iter(qn('w:p')):
        # walk children: chunks of consecutive text runs, split at math
        chunks = []           # list of dicts {runs:[], prev_math:bool, next_math:bool}
        cur = None
        prev_was_math = False
        for child in p_el:
            tag = child.tag
            if tag == qn('w:pPr'):
                continue
            if tag == f'{{{M_NS}}}oMath' or tag == f'{{{M_NS}}}oMathPara':
                if cur is not None:
                    cur['next_math'] = True
                    chunks.append(cur)
                    cur = None
                prev_was_math = True
                continue
            if tag == qn('w:r'):
                ts = child.findall(qn('w:t'))
                if ts:
                    if cur is None:
                        cur = {'runs': [], 'prev_math': prev_was_math, 'next_math': False}
                    cur['runs'].extend(ts)
            # other children ignored
        if cur is not None:
            chunks.append(cur)

        for ci, ch in enumerate(chunks):
            zh = ''.join(t.text or '' for t in ch['runs'])
            if not has_han(zh):
                continue
            key = None
            # find segment key by matching zh text against cache
            for k, v in trans.items():
                if v['zh'] == zh and k.endswith(f':{ci}') and k.rsplit(":",1)[0] == str(_pidx(p_el, body)):
                    pass
            # simpler: use paragraph index
            seg_key = f'{pid}:{ci}'
            entry = trans.get(seg_key)
            if not entry or entry['zh'] != zh:
                continue
            en = mech_fix(entry['en'])
            # restore math-adjacent whitespace to match original chunk boundaries
            if ch['prev_math'] and zh.startswith(' ') and not en.startswith(' '):
                en = ' ' + en
            if ch['next_math'] and zh.endswith(' ') and not en.endswith(' '):
                en = en + ' '
            ts = ch['runs']
            target = max(ts, key=lambda t: len(t.text or ''))
            for t in ts:
                if t is target:
                    t.text = en
                    t.set(qn('xml:space'), 'preserve')
                else:
                    t.text = ''
            applied += 1
    doc.save(DST)
    print(f'saved {DST}; applied {applied}')

_pid_map = {}
def _pidx(p_el, body):
    # paragraph ordinal within body (same as docwalk collect_segments)
    if p_el not in _pid_map:
        for i, p in enumerate(body.iter(qn('w:p'))):
            _pid_map[p] = i
    return _pid_map[p_el]

# need pid available in main loop - restructure: precompute map
def main2():
    trans = json.load(open(TRANS, encoding='utf-8'))
    doc = Document(SRC)
    body = doc.element.body
    pid_map = {p: i for i, p in enumerate(body.iter(qn('w:p')))}
    applied = 0
    for p_el, pid in pid_map.items():
        chunks = []
        cur = None
        prev_was_math = False
        for child in p_el:
            tag = child.tag
            if tag == qn('w:pPr'):
                continue
            if tag == f'{{{M_NS}}}oMath' or tag == f'{{{M_NS}}}oMathPara':
                if cur is not None:
                    cur['next_math'] = True
                    chunks.append(cur)
                    cur = None
                prev_was_math = True
                continue
            if tag == qn('w:r'):
                ts = child.findall(qn('w:t'))
                if ts:
                    if cur is None:
                        cur = {'runs': [], 'prev_math': prev_was_math, 'next_math': False}
                    cur['runs'].extend(ts)
        if cur is not None:
            chunks.append(cur)

        for ci, ch in enumerate(chunks):
            zh = ''.join(t.text or '' for t in ch['runs'])
            if not has_han(zh):
                # punctuation-only chunks (e.g. between math): convert full-width punct
                for t in ch['runs']:
                    if t.text:
                        t.text = (t.text
                                  .replace('、', ', ').replace('，', ', ').replace('。', '. ')
                                  .replace('；', '; ').replace('：', ': ')
                                  .replace('（', '(').replace('）', ')')
                                  .replace('“', '"').replace('”', '"')
                                  .replace('‘', "'").replace('’', "'"))
                continue
            entry = trans.get(f'{pid}:{ci}')
            if not entry or entry['zh'] != zh:
                print(f'WARN miss p{pid}c{ci}: {zh[:40]}')
                continue
            en = mech_fix(entry['en'])
            if ch['prev_math'] and zh.startswith(' ') and not en.startswith(' '):
                en = ' ' + en
            if ch['next_math'] and zh.endswith(' ') and not en.endswith(' '):
                en = en + ' '
            ts = ch['runs']
            target = max(ts, key=lambda t: len(t.text or ''))
            for t in ts:
                if t is target:
                    t.text = en
                    t.set(qn('xml:space'), 'preserve')
                else:
                    t.text = ''
            applied += 1

        # bold label parity with 二稿
        full = ''.join(t.text or '' for t in p_el.iter(qn('w:t')))
        for label in ('Keywords', 'Section summary:', 'Summary of the section:'):
            if full.startswith(label):
                if split_bold_label(p_el, label):
                    print(f'bold label applied: {label}')
                break
    beautify_tables(doc)
    doc.save(DST)
    print(f'saved {DST}; applied {applied}')

if __name__ == '__main__':
    main2()
